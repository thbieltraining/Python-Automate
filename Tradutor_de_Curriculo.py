import os
import time

from deep_translator import MyMemoryTranslator
from docx import Document

# ==========================================
# CONFIGURAÇÃO
# ==========================================

CAMINHO_ENTRADA = r"C:\temp\curriculo\SEU_CURRICULO_2026_CONVERTIDO.docx"
CAMINHO_SAIDA = r"C:\temp\curriculo\SEU_CURRICULO_2026_EN.docx"

LIMITE_TRADUCAO = 450


# ==========================================
# DIVIDIR TEXTO EM BLOCOS
# ==========================================


def dividir_texto(texto, limite=LIMITE_TRADUCAO):

    palavras = texto.split()
    blocos = []
    bloco_atual = ""

    for palavra in palavras:

        candidato = f"{bloco_atual} {palavra}".strip()

        if len(candidato) <= limite:

            bloco_atual = candidato

        else:

            if bloco_atual:
                blocos.append(bloco_atual)

            bloco_atual = palavra

    if bloco_atual:
        blocos.append(bloco_atual)

    return blocos


# ==========================================
# TRADUZIR TEXTO
# ==========================================


def traduzir_texto(texto, translator):

    if not texto or not texto.strip():
        return texto

    blocos = dividir_texto(texto)

    traducoes = []

    print(f"   📦 {len(blocos)} bloco(s) para traduzir")

    for numero, bloco in enumerate(blocos, start=1):

        print(f"   🌎 Bloco {numero}/{len(blocos)} " f"({len(bloco)} caracteres)")

        try:

            traducao = translator.translate(bloco)

            if traducao:

                traducoes.append(traducao)

            else:

                print("   ⚠ Nenhuma tradução retornada.")

                traducoes.append(bloco)

        except (ConnectionError, TimeoutError) as erro:

            print(f"   ⚠ Erro de conexão: {erro}")

            traducoes.append(bloco)

        except ValueError as erro:

            print(f"   ⚠ Erro no texto: {erro}")

            traducoes.append(bloco)

        # Pequena pausa entre requisições
        time.sleep(0.5)

    return " ".join(traducoes)


# ==========================================
# TRADUZIR PARÁGRAFO
# ==========================================


def traduzir_paragrafo(paragrafo, translator):

    texto = paragrafo.text

    if not texto.strip():
        return

    print(f"\n📝 Texto com {len(texto)} caracteres")

    texto_traduzido = traduzir_texto(texto, translator)

    if paragrafo.runs:

        paragrafo.runs[0].text = texto_traduzido

        for run in paragrafo.runs[1:]:

            run.text = ""

    else:

        paragrafo.text = texto_traduzido


# ==========================================
# TRADUZIR TABELAS
# ==========================================


def traduzir_tabelas(doc, translator):

    print("\n" + "=" * 60)
    print("📊 VERIFICANDO TABELAS")
    print("=" * 60)

    if not doc.tables:

        print("Nenhuma tabela encontrada.")

        return

    for numero, tabela in enumerate(doc.tables, start=1):

        print(f"\n📊 Tabela {numero}")

        for linha in tabela.rows:

            for celula in linha.cells:

                for paragrafo in celula.paragraphs:

                    traduzir_paragrafo(paragrafo, translator)


# ==========================================
# FUNÇÃO PRINCIPAL
# ==========================================


def traduzir_curriculo():

    print("=" * 60)
    print("🌎 TRADUTOR DE CURRÍCULO")
    print("Português → Inglês")
    print("=" * 60)

    # ==========================================
    # VERIFICAR ARQUIVO
    # ==========================================

    if not os.path.isfile(CAMINHO_ENTRADA):

        print("\n❌ ARQUIVO NÃO ENCONTRADO!")
        print(CAMINHO_ENTRADA)

        return

    print("\n✅ Arquivo encontrado!")
    print(f"📄 {CAMINHO_ENTRADA}")

    # ==========================================
    # ABRIR DOCUMENTO
    # ==========================================

    try:

        doc = Document(CAMINHO_ENTRADA)

    except (OSError, ValueError) as erro:

        print("\n❌ Erro ao abrir documento:")
        print(erro)

        return

    print("✅ Documento aberto!")

    # ==========================================
    # CONFIGURAR TRADUTOR
    # ==========================================

    translator = MyMemoryTranslator(source="pt-BR", target="en-US")

    print("🌎 MyMemory configurado!")
    print("PT-BR → EN-US")

    # ==========================================
    # TRADUZIR PARÁGRAFOS
    # ==========================================

    print("\n" + "=" * 60)
    print("📝 TRADUZINDO PARÁGRAFOS")
    print("=" * 60)

    total = len(doc.paragraphs)

    print(f"Total de parágrafos: {total}")

    for numero, paragrafo in enumerate(doc.paragraphs, start=1):

        if paragrafo.text.strip():

            print(f"\n[{numero}/{total}]")

            traduzir_paragrafo(paragrafo, translator)

    # ==========================================
    # TRADUZIR TABELAS
    # ==========================================

    traduzir_tabelas(doc, translator)

    # ==========================================
    # SALVAR
    # ==========================================

    print("\n" + "=" * 60)
    print("💾 SALVANDO DOCUMENTO")
    print("=" * 60)

    try:

        doc.save(CAMINHO_SAIDA)

    except (OSError, PermissionError) as erro:

        print("\n❌ Erro ao salvar:")
        print(erro)

        return

    print("\n🎉 TRADUÇÃO CONCLUÍDA!")
    print("=" * 60)
    print(f"\n📁 Arquivo criado:")
    print(CAMINHO_SAIDA)
    print("\n" + "=" * 60)


# ==========================================
# EXECUÇÃO
# ==========================================

if __name__ == "__main__":
    traduzir_curriculo()
